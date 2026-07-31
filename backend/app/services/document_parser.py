"""PDF、DOCX、Markdown 与 TXT 的统一解析及引用锚点。

重依赖全部延迟导入：PDF 优先用 PyMuPDF 获得页面坐标，并在可用时执行 OCR；
DOCX 优先用 python-docx；两者缺失时分别降级到 pypdf/Docling 和标准库 XML。
解析结果只包含可序列化值，可以直接写入 PostgreSQL JSON 字段。
"""

from __future__ import annotations

import hashlib
import math
import mimetypes
import re
import unicodedata
import zipfile
from collections.abc import Callable, Mapping, Sequence
from dataclasses import asdict, dataclass, field, replace
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".md", ".markdown", ".txt"})


class DocumentParseError(RuntimeError):
    """文档无法安全、完整地解析。"""


class UnsupportedDocumentError(DocumentParseError):
    """文件类型不在首版支持范围。"""


@dataclass(frozen=True, slots=True)
class ParserLimits:
    max_file_bytes: int = 200 * 1024 * 1024
    max_pages: int = 2000
    chunk_chars: int = 1400
    chunk_overlap_chars: int = 120
    min_pdf_text_chars_before_ocr: int = 12


@dataclass(frozen=True, slots=True)
class Rectangle:
    x0: float
    y0: float
    x1: float
    y1: float

    def __post_init__(self) -> None:
        coordinates = (self.x0, self.y0, self.x1, self.y1)
        if not all(math.isfinite(value) for value in coordinates):
            raise ValueError("矩形坐标必须是有限数值")
        left, right = sorted((self.x0, self.x1))
        top, bottom = sorted((self.y0, self.y1))
        object.__setattr__(self, "x0", left)
        object.__setattr__(self, "y0", top)
        object.__setattr__(self, "x1", right)
        object.__setattr__(self, "y1", bottom)


@dataclass(frozen=True, slots=True)
class CitationAnchor:
    document_id: str
    document_version_id: str
    block_id: str
    quote: str
    quote_hash: str
    page_number: int | None = None  # 面向用户，始终从 1 开始
    page_width: float | None = None
    page_height: float | None = None
    rects: tuple[Rectangle, ...] = ()
    char_start: int | None = None
    char_end: int | None = None
    coordinate_system: str | None = None


@dataclass(frozen=True, slots=True)
class DocumentAsset:
    asset_id: str
    kind: str
    media_type: str | None
    content_hash: str
    storage_path: str | None = None
    page_number: int | None = None
    rect: Rectangle | None = None
    metadata: Mapping[str, Any] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class ParsedBlock:
    block_id: str
    kind: str
    text: str
    section_path: tuple[str, ...]
    anchor: CitationAnchor
    metadata: Mapping[str, Any] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class ParsedChunk:
    chunk_id: str
    text: str
    section_path: tuple[str, ...]
    anchors: tuple[CitationAnchor, ...]
    block_ids: tuple[str, ...]
    token_estimate: int
    metadata: Mapping[str, Any] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    document_id: str
    document_version_id: str
    filename: str
    extension: str
    mime_type: str
    content_hash: str
    parser_backend: str
    blocks: tuple[ParsedBlock, ...]
    chunks: tuple[ParsedChunk, ...]
    assets: tuple[DocumentAsset, ...] = ()
    warnings: tuple[str, ...] = ()
    metadata: Mapping[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _stable_id(prefix: str, *parts: Any) -> str:
    payload = "\x1f".join(str(part) for part in parts).encode("utf-8")
    return f"{prefix}_{hashlib.sha256(payload).hexdigest()[:24]}"


def _quote_hash(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", text).strip()
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for part in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(part)
    return digest.hexdigest()


def detect_text_encoding(data: bytes) -> tuple[str, str | None]:
    """返回可解码中文文本的编码与可选警告。"""

    if data.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig", None
    if data.startswith((b"\xff\xfe", b"\xfe\xff")):
        return "utf-16", None
    try:
        data.decode("utf-8")
        return "utf-8", None
    except UnicodeDecodeError:
        pass
    try:
        from charset_normalizer import from_bytes

        match = from_bytes(data).best()
        if match and match.encoding:
            return match.encoding, f"文本编码识别为 {match.encoding}"
    except ImportError:
        pass
    for encoding in ("gb18030", "big5"):
        try:
            data.decode(encoding)
            return encoding, f"文本按 {encoding} 解码"
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("无法识别文本编码；建议转换为 UTF-8")


class UnifiedDocumentParser:
    def __init__(
        self,
        *,
        limits: ParserLimits | None = None,
        ocr_languages: str = "chi_sim+eng",
        asset_dir: str | Path | None = None,
        progress_callback: Callable[[int, int, str], None] | None = None,
    ) -> None:
        self.limits = limits or ParserLimits()
        self.ocr_languages = ocr_languages
        self.asset_dir = Path(asset_dir) if asset_dir else None
        self.progress_callback = progress_callback

    def _notify_progress(self, current: int, total: int, stage: str) -> None:
        if self.progress_callback:
            self.progress_callback(current, max(1, total), stage)

    def parse(
        self,
        path: str | Path,
        *,
        document_id: str = "document",
        document_version_id: str = "version",
    ) -> ParsedDocument:
        source = Path(path)
        if not source.is_file():
            raise DocumentParseError(f"文件不存在：{source.name}")
        extension = source.suffix.casefold()
        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentError(f"暂不支持 {extension or '无扩展名'} 文件")
        size = source.stat().st_size
        if size > self.limits.max_file_bytes:
            raise DocumentParseError(f"文件超过 {self.limits.max_file_bytes} 字节限制")
        content_hash = _file_hash(source)

        if extension == ".pdf":
            parsed = self._parse_pdf(source, document_id, document_version_id, content_hash)
        elif extension == ".docx":
            parsed = self._parse_docx(source, document_id, document_version_id, content_hash)
        else:
            raw = source.read_bytes()
            encoding, warning = detect_text_encoding(raw)
            # 非 PDF 文本锚点统一基于 LF；浏览器、编辑器与 JSON 均可稳定复现该位置。
            text = raw.decode(encoding).replace("\r\n", "\n").replace("\r", "\n")
            blocks = (
                self._markdown_blocks(text, document_id, document_version_id)
                if extension in {".md", ".markdown"}
                else self._text_blocks(text, document_id, document_version_id)
            )
            self._notify_progress(1, 1, "text_structure")
            mime = "text/markdown" if extension in {".md", ".markdown"} else "text/plain"
            parsed = self._build_document(
                source,
                document_id,
                document_version_id,
                content_hash,
                mime,
                f"native-{encoding}",
                blocks,
                warnings=(warning,) if warning else (),
            )
        return parsed

    def _build_document(
        self,
        source: Path,
        document_id: str,
        version_id: str,
        content_hash: str,
        mime_type: str,
        backend: str,
        blocks: Sequence[ParsedBlock],
        *,
        assets: Sequence[DocumentAsset] = (),
        warnings: Sequence[str] = (),
        metadata: Mapping[str, Any] | None = None,
    ) -> ParsedDocument:
        chunks = self._make_chunks(blocks, document_id, version_id)
        return ParsedDocument(
            document_id=str(document_id),
            document_version_id=str(version_id),
            filename=source.name,
            extension=source.suffix.casefold(),
            mime_type=mime_type,
            content_hash=content_hash,
            parser_backend=backend,
            blocks=tuple(blocks),
            chunks=tuple(chunks),
            assets=tuple(assets),
            warnings=tuple(dict.fromkeys(warnings)),
            metadata=dict(metadata or {}),
        )

    def _parse_pdf(
        self, source: Path, document_id: str, version_id: str, content_hash: str
    ) -> ParsedDocument:
        try:
            return self._parse_pdf_pymupdf(source, document_id, version_id, content_hash)
        except ImportError:
            pass
        try:
            return self._parse_pdf_pypdf(source, document_id, version_id, content_hash)
        except ImportError:
            pass
        docling = self._parse_with_docling(source, document_id, version_id, content_hash)
        if docling is not None:
            return replace(docling, warnings=docling.warnings + ("PDF 坐标不可用：未安装 PyMuPDF",))
        raise DocumentParseError("解析 PDF 需要安装 PyMuPDF、pypdf 或 Docling")

    def _parse_pdf_pymupdf(
        self, source: Path, document_id: str, version_id: str, content_hash: str
    ) -> ParsedDocument:
        import fitz  # PyMuPDF

        warnings: list[str] = []
        blocks: list[ParsedBlock] = []
        assets: list[DocumentAsset] = []
        pages_meta: list[dict[str, Any]] = []
        with fitz.open(source) as pdf:
            if pdf.page_count > self.limits.max_pages:
                raise DocumentParseError(f"PDF 超过 {self.limits.max_pages} 页限制")
            for page_index, page in enumerate(pdf):
                page_number = page_index + 1
                width, height = float(page.rect.width), float(page.rect.height)
                raw = page.get_text("dict", sort=True)
                visible_text = "".join(
                    str(span.get("text", ""))
                    for block in raw.get("blocks", [])
                    if block.get("type") == 0
                    for line in block.get("lines", [])
                    for span in line.get("spans", [])
                ).strip()
                used_ocr = False
                if len(visible_text) < self.limits.min_pdf_text_chars_before_ocr:
                    try:
                        text_page = page.get_textpage_ocr(
                            language=self.ocr_languages,
                            dpi=200,
                            full=True,
                        )
                        raw = page.get_text("dict", sort=True, textpage=text_page)
                        used_ocr = True
                        warnings.append(f"第 {page_number} 页使用 OCR")
                    except Exception:
                        warnings.append(f"第 {page_number} 页文本较少且 OCR 不可用")
                pages_meta.append(
                    {
                        "page_number": page_number,
                        "width": width,
                        "height": height,
                        "used_ocr": used_ocr,
                    }
                )
                for block_index, raw_block in enumerate(raw.get("blocks", [])):
                    block_type = raw_block.get("type")
                    bbox = raw_block.get("bbox") or (0, 0, 0, 0)
                    if block_type == 0:
                        line_texts: list[str] = []
                        rects: list[Rectangle] = []
                        for line in raw_block.get("lines", []):
                            spans = line.get("spans", [])
                            line_text = "".join(str(span.get("text", "")) for span in spans).strip()
                            if line_text:
                                line_texts.append(line_text)
                                line_box = line.get("bbox") or spans[0].get("bbox")
                                if line_box:
                                    rects.append(Rectangle(*(round(float(v), 3) for v in line_box)))
                        text = "\n".join(line_texts).strip()
                        if not text:
                            continue
                        block_id = _stable_id("block", version_id, page_number, block_index, text)
                        anchor = CitationAnchor(
                            str(document_id),
                            str(version_id),
                            block_id,
                            text,
                            _quote_hash(text),
                            page_number=page_number,
                            page_width=width,
                            page_height=height,
                            rects=tuple(rects) or (Rectangle(*(float(v) for v in bbox)),),
                            coordinate_system="pdf_points_top_left",
                        )
                        blocks.append(
                            ParsedBlock(
                                block_id,
                                "text",
                                text,
                                (f"第 {page_number} 页",),
                                anchor,
                                {"ocr": used_ocr},
                            )
                        )
                    elif block_type == 1:
                        image_bytes = raw_block.get("image")
                        if not image_bytes:
                            continue
                        image_hash = hashlib.sha256(image_bytes).hexdigest()
                        extension = str(raw_block.get("ext") or "png").casefold()
                        storage_path = self._store_asset(image_hash, extension, image_bytes)
                        rect = Rectangle(*(round(float(v), 3) for v in bbox))
                        asset_id = _stable_id(
                            "asset",
                            version_id,
                            page_number,
                            block_index,
                            image_hash,
                        )
                        assets.append(
                            DocumentAsset(
                                asset_id,
                                "image",
                                mimetypes.guess_type(f"x.{extension}")[0],
                                image_hash,
                                storage_path,
                                page_number,
                                rect,
                                {
                                    "requires_vision": True,
                                    "width": raw_block.get("width"),
                                    "height": raw_block.get("height"),
                                },
                            )
                        )
                self._extract_pdf_tables(
                    page, document_id, version_id, page_number, width, height, blocks, warnings
                )
                self._notify_progress(page_number, pdf.page_count, "pdf_pages")
        if not any(block.text.strip() for block in blocks):
            docling = self._parse_with_docling(source, document_id, version_id, content_hash)
            if docling is not None:
                return replace(
                    docling,
                    assets=tuple(assets),
                    warnings=tuple(warnings)
                    + docling.warnings
                    + ("OCR 结果无坐标，出处只能定位到结构化段落",),
                    metadata={**docling.metadata, "pages": pages_meta},
                )
        return self._build_document(
            source,
            document_id,
            version_id,
            content_hash,
            "application/pdf",
            "pymupdf",
            blocks,
            assets=assets,
            warnings=warnings,
            metadata={"page_count": len(pages_meta), "pages": pages_meta},
        )

    def _extract_pdf_tables(
        self,
        page: Any,
        document_id: str,
        version_id: str,
        page_number: int,
        width: float,
        height: float,
        blocks: list[ParsedBlock],
        warnings: list[str],
    ) -> None:
        if not hasattr(page, "find_tables"):
            return
        try:
            tables = page.find_tables()
            for table_index, table in enumerate(tables.tables):
                rows = table.extract()
                text = "\n".join(
                    "\t".join(str(cell or "").strip() for cell in row) for row in rows
                ).strip()
                if not text:
                    continue
                block_id = _stable_id("table", version_id, page_number, table_index, text)
                bbox = Rectangle(*(round(float(v), 3) for v in table.bbox))
                anchor = CitationAnchor(
                    str(document_id),
                    str(version_id),
                    block_id,
                    text,
                    _quote_hash(text),
                    page_number,
                    width,
                    height,
                    (bbox,),
                    coordinate_system="pdf_points_top_left",
                )
                blocks.append(
                    ParsedBlock(block_id, "table", text, (f"第 {page_number} 页",), anchor)
                )
        except Exception:
            warnings.append(f"第 {page_number} 页表格结构识别失败，已保留普通文本")

    def _parse_pdf_pypdf(
        self, source: Path, document_id: str, version_id: str, content_hash: str
    ) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(source))
        if len(reader.pages) > self.limits.max_pages:
            raise DocumentParseError(f"PDF 超过 {self.limits.max_pages} 页限制")
        blocks: list[ParsedBlock] = []
        for page_index, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            page_number = page_index + 1
            block_id = _stable_id("block", version_id, page_number, text)
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            anchor = CitationAnchor(
                str(document_id),
                str(version_id),
                block_id,
                text,
                _quote_hash(text),
                page_number,
                width,
                height,
                coordinate_system="pdf_points_bottom_left",
            )
            blocks.append(ParsedBlock(block_id, "text", text, (f"第 {page_number} 页",), anchor))
            self._notify_progress(page_number, len(reader.pages), "pdf_pages")
        return self._build_document(
            source,
            document_id,
            version_id,
            content_hash,
            "application/pdf",
            "pypdf",
            blocks,
            warnings=("未安装 PyMuPDF，PDF 出处只有页码、没有高亮矩形",),
            metadata={"page_count": len(reader.pages)},
        )

    def _parse_docx(
        self, source: Path, document_id: str, version_id: str, content_hash: str
    ) -> ParsedDocument:
        try:
            blocks = self._parse_docx_python(source, document_id, version_id)
            backend = "python-docx"
        except ImportError:
            blocks = self._parse_docx_xml(source, document_id, version_id)
            backend = "stdlib-docx-xml"
        assets = self._docx_assets(source, version_id)
        self._notify_progress(1, 1, "docx_structure")
        return self._build_document(
            source,
            document_id,
            version_id,
            content_hash,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            backend,
            blocks,
            assets=assets,
            warnings=("DOCX 没有稳定页码，出处按段落定位",),
        )

    def _parse_docx_python(
        self, source: Path, document_id: str, version_id: str
    ) -> list[ParsedBlock]:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(str(source))
        blocks: list[ParsedBlock] = []
        section: list[str] = []
        cursor = 0
        for index, child in enumerate(document.element.body.iterchildren()):
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                if not text:
                    continue
                style_name = str(paragraph.style.name if paragraph.style else "")
                heading = re.search(r"(?:heading|标题)\s*([1-6])?", style_name, re.IGNORECASE)
                kind = "heading" if heading else "text"
                if heading:
                    level = int(heading.group(1) or 1)
                    section[level - 1 :] = [text]
            elif tag == "tbl":
                table = Table(child, document)
                text = "\n".join(
                    "\t".join(cell.text.strip() for cell in row.cells) for row in table.rows
                ).strip()
                kind = "table"
                if not text:
                    continue
            else:
                continue
            block_id = _stable_id("block", version_id, index, text)
            anchor = CitationAnchor(
                str(document_id),
                str(version_id),
                block_id,
                text,
                _quote_hash(text),
                char_start=cursor,
                char_end=cursor + len(text),
            )
            blocks.append(ParsedBlock(block_id, kind, text, tuple(section), anchor))
            cursor += len(text) + 2
        return blocks

    def _parse_docx_xml(self, source: Path, document_id: str, version_id: str) -> list[ParsedBlock]:
        try:
            with zipfile.ZipFile(source) as archive:
                xml = archive.read("word/document.xml")
        except (zipfile.BadZipFile, KeyError) as exc:
            raise DocumentParseError("DOCX 文件结构损坏") from exc
        root = ElementTree.fromstring(xml)
        body = next((node for node in root.iter() if node.tag.rsplit("}", 1)[-1] == "body"), None)
        if body is None:
            return []
        blocks: list[ParsedBlock] = []
        section: list[str] = []
        cursor = 0
        for index, node in enumerate(body):
            tag = node.tag.rsplit("}", 1)[-1]
            if tag == "p":
                text = "".join(
                    (item.text or "") for item in node.iter() if item.tag.rsplit("}", 1)[-1] == "t"
                ).strip()
                style = next(
                    (
                        str(item.attrib.get(next(iter(item.attrib), ""), ""))
                        for item in node.iter()
                        if item.tag.rsplit("}", 1)[-1] == "pStyle"
                    ),
                    "",
                )
                heading = re.search(r"(?:heading|标题)([1-6])?", style, re.IGNORECASE)
                kind = "heading" if heading else "text"
                if heading:
                    level = int(heading.group(1) or 1)
                    section[level - 1 :] = [text]
            elif tag == "tbl":
                rows: list[str] = []
                for row in (item for item in node.iter() if item.tag.rsplit("}", 1)[-1] == "tr"):
                    cells = []
                    for cell in (item for item in row if item.tag.rsplit("}", 1)[-1] == "tc"):
                        cells.append(
                            "".join(
                                (part.text or "")
                                for part in cell.iter()
                                if part.tag.rsplit("}", 1)[-1] == "t"
                            )
                        )
                    rows.append("\t".join(cells))
                text, kind = "\n".join(rows).strip(), "table"
            else:
                continue
            if not text:
                continue
            block_id = _stable_id("block", version_id, index, text)
            anchor = CitationAnchor(
                str(document_id),
                str(version_id),
                block_id,
                text,
                _quote_hash(text),
                char_start=cursor,
                char_end=cursor + len(text),
            )
            blocks.append(ParsedBlock(block_id, kind, text, tuple(section), anchor))
            cursor += len(text) + 2
        return blocks

    def _docx_assets(self, source: Path, version_id: str) -> list[DocumentAsset]:
        assets: list[DocumentAsset] = []
        try:
            with zipfile.ZipFile(source) as archive:
                for name in archive.namelist():
                    if not name.casefold().startswith("word/media/") or name.endswith("/"):
                        continue
                    data = archive.read(name)
                    digest = hashlib.sha256(data).hexdigest()
                    extension = Path(name).suffix.lstrip(".") or "bin"
                    assets.append(
                        DocumentAsset(
                            _stable_id("asset", version_id, digest),
                            "image",
                            mimetypes.guess_type(name)[0],
                            digest,
                            self._store_asset(digest, extension, data),
                            metadata={"source_name": Path(name).name, "requires_vision": True},
                        )
                    )
        except zipfile.BadZipFile:
            pass
        return assets

    def _store_asset(self, digest: str, extension: str, data: bytes) -> str | None:
        if self.asset_dir is None:
            return None
        self.asset_dir.mkdir(parents=True, exist_ok=True)
        target = self.asset_dir / f"{digest}.{re.sub(r'[^a-zA-Z0-9]', '', extension) or 'bin'}"
        if not target.exists():
            target.write_bytes(data)
        return str(target)

    def _parse_with_docling(
        self, source: Path, document_id: str, version_id: str, content_hash: str
    ) -> ParsedDocument | None:
        try:
            from docling.document_converter import DocumentConverter
        except ImportError:
            return None
        try:
            result = DocumentConverter().convert(str(source))
            markdown = result.document.export_to_markdown()
            blocks = self._markdown_blocks(markdown, document_id, version_id)
            self._notify_progress(1, 1, "docling")
        except Exception as exc:
            raise DocumentParseError(f"Docling 解析失败：{type(exc).__name__}") from exc
        return self._build_document(
            source,
            document_id,
            version_id,
            content_hash,
            mimetypes.guess_type(source.name)[0] or "application/octet-stream",
            "docling",
            blocks,
            warnings=("Docling 结构化结果未提供稳定原文件坐标",),
        )

    def _markdown_blocks(self, text: str, document_id: str, version_id: str) -> list[ParsedBlock]:
        lines = text.splitlines(keepends=True)
        offsets: list[int] = []
        cursor = 0
        for line in lines:
            offsets.append(cursor)
            cursor += len(line)
        blocks: list[ParsedBlock] = []
        section: list[str] = []
        index = 0
        while index < len(lines):
            raw = lines[index]
            stripped = raw.strip()
            if not stripped:
                index += 1
                continue
            start_line = index
            heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
            if heading:
                level, content = len(heading.group(1)), heading.group(2)
                section[level - 1 :] = [content]
                kind, value = "heading", content
                index += 1
            elif stripped.startswith("```") or stripped.startswith("~~~"):
                marker = stripped[:3]
                index += 1
                while index < len(lines) and not lines[index].strip().startswith(marker):
                    index += 1
                if index < len(lines):
                    index += 1
                kind = "code"
                value = "".join(lines[start_line:index]).strip()
            elif (
                "|" in stripped
                and index + 1 < len(lines)
                and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1])
            ):
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    index += 1
                kind = "table"
                value = "".join(lines[start_line:index]).strip()
            else:
                index += 1
                while index < len(lines):
                    candidate = lines[index]
                    if not candidate.strip() or re.match(r"^#{1,6}\s+", candidate.strip()):
                        break
                    if candidate.strip().startswith(("```", "~~~")):
                        break
                    index += 1
                kind = "text"
                value = "".join(lines[start_line:index]).strip()
            end = offsets[index] if index < len(offsets) else len(text)
            start = offsets[start_line]
            block_id = _stable_id("block", version_id, start, value)
            anchor = CitationAnchor(
                str(document_id),
                str(version_id),
                block_id,
                value,
                _quote_hash(value),
                char_start=start,
                char_end=end,
            )
            blocks.append(ParsedBlock(block_id, kind, value, tuple(section), anchor))
        return blocks

    def _text_blocks(self, text: str, document_id: str, version_id: str) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        for index, match in enumerate(
            re.finditer(r"\S(?:.*?\S)?(?=\r?\n\s*\r?\n|\Z)", text, re.DOTALL)
        ):
            value = match.group(0).strip()
            if not value:
                continue
            start, end = match.start(), match.end()
            block_id = _stable_id("block", version_id, start, value)
            anchor = CitationAnchor(
                str(document_id),
                str(version_id),
                block_id,
                value,
                _quote_hash(value),
                char_start=start,
                char_end=end,
            )
            blocks.append(
                ParsedBlock(block_id, "text", value, (), anchor, {"paragraph": index + 1})
            )
        return blocks

    def _make_chunks(
        self, blocks: Sequence[ParsedBlock], document_id: str, version_id: str
    ) -> list[ParsedChunk]:
        max_chars = max(200, self.limits.chunk_chars)
        chunks: list[ParsedChunk] = []
        group: list[ParsedBlock] = []
        group_chars = 0

        def flush() -> None:
            nonlocal group, group_chars
            if not group:
                return
            content = "\n\n".join(item.text for item in group).strip()
            if content:
                chunk_id = _stable_id(
                    "chunk", version_id, *(item.block_id for item in group), content
                )
                chunks.append(
                    ParsedChunk(
                        chunk_id,
                        content,
                        group[-1].section_path,
                        tuple(item.anchor for item in group),
                        tuple(item.block_id for item in group),
                        max(1, (len(content) + 2) // 3),
                        {"document_id": str(document_id), "document_version_id": str(version_id)},
                    )
                )
            group, group_chars = [], 0

        for block in blocks:
            text = block.text.strip()
            if not text:
                continue
            if len(text) > max_chars:
                flush()
                step = max(1, max_chars - min(self.limits.chunk_overlap_chars, max_chars // 4))
                for start in range(0, len(text), step):
                    part = text[start : start + max_chars].strip()
                    if not part:
                        continue
                    chunk_id = _stable_id("chunk", version_id, block.block_id, start, part)
                    chunks.append(
                        ParsedChunk(
                            chunk_id,
                            part,
                            block.section_path,
                            (block.anchor,),
                            (block.block_id,),
                            max(1, (len(part) + 2) // 3),
                            {
                                "document_id": str(document_id),
                                "document_version_id": str(version_id),
                                "part_start": start,
                            },
                        )
                    )
                    if start + max_chars >= len(text):
                        break
                continue
            same_scope = not group or (
                group[-1].anchor.page_number == block.anchor.page_number
                and group[-1].section_path == block.section_path
            )
            if group and (not same_scope or group_chars + len(text) + 2 > max_chars):
                flush()
            group.append(block)
            group_chars += len(text) + 2
        flush()
        return chunks


def parse_document(
    path: str | Path,
    *,
    document_id: str = "document",
    document_version_id: str = "version",
    limits: ParserLimits | None = None,
    asset_dir: str | Path | None = None,
) -> ParsedDocument:
    """便捷入口，适合 CLI 和 Celery 调用。"""

    return UnifiedDocumentParser(limits=limits, asset_dir=asset_dir).parse(
        path,
        document_id=document_id,
        document_version_id=document_version_id,
    )
