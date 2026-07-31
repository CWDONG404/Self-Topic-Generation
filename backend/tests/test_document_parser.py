from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

from app.services.document_parser import Rectangle, UnifiedDocumentParser


def test_rectangle_normalizes_rotated_pdf_coordinates() -> None:
    rectangle = Rectangle(180.0, 220.0, 40.0, 80.0)

    assert rectangle == Rectangle(40.0, 80.0, 180.0, 220.0)


def test_markdown_keeps_sections_and_character_anchors(tmp_path: Path) -> None:
    source = tmp_path / "考点.md"
    source.write_text(
        "# 计算环境安全\n\n访问控制用于限制主体对客体的访问。\n\n"
        "## 审计\n\n审计日志应受到完整性保护。\n",
        encoding="utf-8",
    )
    parsed = UnifiedDocumentParser().parse(source, document_id="doc", document_version_id="v1")
    assert parsed.parser_backend == "native-utf-8"
    assert parsed.blocks
    assert parsed.chunks
    audit = next(block for block in parsed.blocks if "审计日志" in block.text)
    assert audit.section_path == ("计算环境安全", "审计")
    assert audit.anchor.char_start is not None
    assert audit.anchor.char_end is not None
    assert (
        source.read_text(encoding="utf-8")[audit.anchor.char_start : audit.anchor.char_end].strip()
        == audit.text
    )


def test_txt_supports_common_chinese_encoding(tmp_path: Path) -> None:
    source = tmp_path / "资料.txt"
    source.write_bytes("第一章\r\n\r\n信息安全管理制度应定期评审。".encode("gb18030"))
    parsed = UnifiedDocumentParser().parse(source, document_id="doc", document_version_id="v2")
    assert "信息安全管理制度" in "\n".join(block.text for block in parsed.blocks)
    assert all(chunk.anchors for chunk in parsed.chunks)


def test_docx_paragraphs_and_tables_are_structured(tmp_path: Path) -> None:
    source = tmp_path / "材料.docx"
    try:
        from docx import Document
    except ImportError:
        document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
        <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:body>
            <w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>安全运营</w:t></w:r></w:p>
            <w:p><w:r><w:t>安全事件应按照预案分级响应。</w:t></w:r></w:p>
            <w:tbl>
              <w:tr><w:tc><w:p><w:r><w:t>等级</w:t></w:r></w:p></w:tc><w:tc><w:p><w:r><w:t>要求</w:t></w:r></w:p></w:tc></w:tr>
              <w:tr><w:tc><w:p><w:r><w:t>一级</w:t></w:r></w:p></w:tc><w:tc><w:p><w:r><w:t>立即响应</w:t></w:r></w:p></w:tc></w:tr>
            </w:tbl>
          </w:body>
        </w:document>"""
        with zipfile.ZipFile(source, "w") as archive:
            archive.writestr("word/document.xml", document_xml)
    else:
        document = Document()
        document.add_heading("安全运营", level=1)
        document.add_paragraph("安全事件应按照预案分级响应。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "等级"
        table.cell(0, 1).text = "要求"
        table.cell(1, 0).text = "一级"
        table.cell(1, 1).text = "立即响应"
        document.save(source)

    parsed = UnifiedDocumentParser().parse(source, document_id="doc", document_version_id="v3")
    assert any(block.kind == "heading" for block in parsed.blocks)
    assert any(block.kind == "table" and "立即响应" in block.text for block in parsed.blocks)
    assert parsed.assets == ()


def test_pdf_has_one_based_page_and_highlight_rectangles(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    source = tmp_path / "坐标.pdf"
    pdf = fitz.open()
    page = pdf.new_page(width=400, height=300)
    page.insert_text((40, 80), "Access control protects resources.")
    pdf.save(source)
    pdf.close()

    parsed = UnifiedDocumentParser().parse(source, document_id="doc", document_version_id="v4")
    block = next(block for block in parsed.blocks if "Access control" in block.text)
    assert block.anchor.page_number == 1
    assert block.anchor.page_width == pytest.approx(400)
    assert block.anchor.page_height == pytest.approx(300)
    assert block.anchor.rects
    assert block.anchor.coordinate_system == "pdf_points_top_left"


def test_pdf_repeated_image_has_unique_asset_instance_ids(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    source = tmp_path / "重复图片.pdf"
    pixel = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 2, 2), False)
    pixel.clear_with(255)
    pdf = fitz.open()
    for page_number in (1, 2):
        page = pdf.new_page(width=300, height=200)
        page.insert_text((40, 100), f"Page {page_number} evidence.")
        page.insert_image(fitz.Rect(20, 20, 60, 60), pixmap=pixel)
    pdf.save(source)
    pdf.close()

    parsed = UnifiedDocumentParser().parse(
        source,
        document_id="doc",
        document_version_id="repeated-image-v1",
    )

    assert len(parsed.assets) == 2
    assert len({asset.asset_id for asset in parsed.assets}) == 2
    assert len({asset.content_hash for asset in parsed.assets}) == 1
    assert {asset.page_number for asset in parsed.assets} == {1, 2}
